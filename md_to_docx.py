"""Convert the project Markdown files into Word (.docx) documents.

Handles headings, paragraphs, bullet lists, pipe tables, bold/italic
inline styling, and inline code. Not a full Markdown parser — tuned
for the files in this repository.

Run:
    python3 md_to_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
OUT_DIR = HERE / "outputs"
OUT_DIR.mkdir(exist_ok=True)

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?[\s\-:|]+\|?\s*$")
BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
NUMBERED_RE = re.compile(r"^\s*(\d+)\.\s+(.*)$")


def _split_cells(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _add_inline(paragraph, text: str):
    """Render **bold**, *italic*, and `code` inline runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*"):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x66, 0x00, 0x99)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            text = row[j] if j < len(row) else ""
            if i == 0:
                run = p.add_run(text)
                run.bold = True
            else:
                _add_inline(p, text)


def _set_base_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def convert(md_path: Path, docx_path: Path, title: str | None = None):
    text = md_path.read_text()
    lines = text.splitlines()

    doc = Document()
    _set_base_style(doc)
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    if title:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(title)
        run.bold = True
        run.font.size = Pt(16)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        m = HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        # Table: header row followed by separator row
        if stripped.startswith("|") and i + 1 < len(lines) \
                and TABLE_SEP_RE.match(lines[i + 1]):
            header = _split_cells(stripped)
            rows = [header]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_split_cells(lines[i]))
                i += 1
            _add_table(doc, rows)
            doc.add_paragraph()
            continue

        # Bulleted list
        if BULLET_RE.match(line):
            while i < len(lines) and BULLET_RE.match(lines[i]):
                content = BULLET_RE.match(lines[i]).group(1)
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, content)
                i += 1
            continue

        # Numbered list
        if NUMBERED_RE.match(line):
            while i < len(lines) and NUMBERED_RE.match(lines[i]):
                content = NUMBERED_RE.match(lines[i]).group(2)
                p = doc.add_paragraph(style="List Number")
                _add_inline(p, content)
                i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped[2:])
            run.italic = True
            i += 1
            continue

        # Plain paragraph (coalesce consecutive non-empty, non-special lines)
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            ns = nxt.strip()
            if not ns:
                break
            if HEADING_RE.match(ns) or ns.startswith("|") \
                    or BULLET_RE.match(nxt) or NUMBERED_RE.match(nxt) \
                    or ns.startswith("> "):
                break
            buf.append(ns)
            i += 1
        para = doc.add_paragraph()
        _add_inline(para, " ".join(buf))

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


def main():
    convert(HERE / "METHODS.md",
            OUT_DIR / "CPE415_Methods.docx",
            title="CPE 415 — Methods")
    convert(HERE / "outputs" / "report.md",
            OUT_DIR / "CPE415_Results.docx",
            title="CPE 415 — Results")
    convert(HERE / "README.md",
            OUT_DIR / "CPE415_README.docx",
            title="CPE 415 — Project Overview")


if __name__ == "__main__":
    main()
